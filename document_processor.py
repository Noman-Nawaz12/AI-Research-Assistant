"""
document_processor.py
----------------------
Handles extraction of text (and basic metadata) from uploaded documents.
Supports: PDF (via PyMuPDF), DOCX (via python-docx), TXT (native Python).
"""

import fitz  # PyMuPDF
import docx  # python-docx
import io


class DocumentProcessingError(Exception):
    """Custom exception for document processing failures."""
    pass


def extract_text_from_pdf(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Extract text from a PDF file, page by page.
    Returns a list of dicts: {"text": ..., "source": filename, "page": page_number}
    """
    chunks = []
    try:
        pdf_stream = io.BytesIO(file_bytes)
        doc = fitz.open(stream=pdf_stream, filetype="pdf")

        if doc.page_count == 0:
            raise DocumentProcessingError(f"'{filename}' appears to be empty (0 pages).")

        for page_num, page in enumerate(doc, start=1):
            text = page.get_text().strip()
            if text:  # skip blank pages
                chunks.append({
                    "text": text,
                    "source": filename,
                    "page": page_num
                })
        doc.close()

        if not chunks:
            raise DocumentProcessingError(f"No extractable text found in '{filename}'. It may be a scanned/image-only PDF.")

        return chunks

    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(f"Failed to process PDF '{filename}': {str(e)}")


def extract_text_from_docx(file_bytes: bytes, filename: str, chars_per_page: int = 1800) -> list[dict]:
    """
    Extract text from a DOCX file, tracking page numbers by:
    1. Detecting explicit page-break markers in the document XML (exact), OR
    2. Falling back to an estimated page number based on character position
       (approximate, since Word does not store natural page breaks in the file).
    Returns a list of dicts: {"text": ..., "source": filename, "page": page_number}
    """
    try:
        doc_stream = io.BytesIO(file_bytes)
        document = docx.Document(doc_stream)

        w_ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

        pages = {}
        current_page = 1
        char_count_on_page = 0
        has_any_explicit_break = False

        # First pass: check if the document has ANY explicit page break markers
        for para in document.paragraphs:
            for run in para.runs:
                for br in run._element.findall(f"{w_ns}br"):
                    if br.get(f"{w_ns}type") == "page":
                        has_any_explicit_break = True

        for para in document.paragraphs:
            para_text = para.text.strip()

            has_page_break = False
            for run in para.runs:
                for br in run._element.findall(f"{w_ns}br"):
                    if br.get(f"{w_ns}type") == "page":
                        has_page_break = True

            if para_text:
                pages.setdefault(current_page, []).append(para_text)
                char_count_on_page += len(para_text)

            if has_page_break:
                current_page += 1
                char_count_on_page = 0
            elif not has_any_explicit_break and char_count_on_page >= chars_per_page:
                # Fallback: no explicit breaks in this doc at all -> estimate by length
                current_page += 1
                char_count_on_page = 0

        if not pages:
            raise DocumentProcessingError(f"No extractable text found in '{filename}'.")

        result = []
        for page_num, paragraphs in pages.items():
            result.append({
                "text": "\n".join(paragraphs),
                "source": filename,
                "page": page_num
            })

        return result

    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(f"Failed to process DOCX '{filename}': {str(e)}")


def extract_text_from_txt(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Extract text from a TXT file.
    Returns a list of dicts: {"text": ..., "source": filename, "page": 1}
    """
    try:
        text = file_bytes.decode("utf-8", errors="ignore").strip()

        if not text:
            raise DocumentProcessingError(f"'{filename}' is empty.")

        return [{
            "text": text,
            "source": filename,
            "page": "N/A"
        }]

    except DocumentProcessingError:
        raise
    except Exception as e:
        raise DocumentProcessingError(f"Failed to process TXT '{filename}': {str(e)}")


def process_document(uploaded_file) -> list[dict]:
    """
    Main entry point: takes a Streamlit UploadedFile object,
    routes it to the correct parser based on extension.
    Returns a list of {"text", "source", "page"} dicts.
    """
    filename = uploaded_file.name
    file_bytes = uploaded_file.read()
    extension = filename.lower().split(".")[-1]

    if extension == "pdf":
        return extract_text_from_pdf(file_bytes, filename)
    elif extension == "docx":
        return extract_text_from_docx(file_bytes, filename)
    elif extension == "txt":
        return extract_text_from_txt(file_bytes, filename)
    else:
        raise DocumentProcessingError(
            f"Unsupported file format: '.{extension}'. Please upload PDF, DOCX, or TXT files only."
        )


def process_multiple_documents(uploaded_files: list) -> tuple[list[dict], list[str]]:
    """
    Process multiple uploaded files.
    Returns (all_extracted_chunks, list_of_error_messages)
    Continues processing other files even if one fails.
    """
    all_data = []
    errors = []

    if not uploaded_files:
        errors.append("No files were uploaded. Please upload at least one document.")
        return all_data, errors

    for uploaded_file in uploaded_files:
        try:
            data = process_document(uploaded_file)
            all_data.extend(data)
        except DocumentProcessingError as e:
            errors.append(str(e))
        except Exception as e:
            errors.append(f"Unexpected error processing '{uploaded_file.name}': {str(e)}")

    return all_data, errors